"""
Genera el documento Word de presentacion del sistema AMUCO para Harold.
Ejecutar: /c/Python312/python.exe docs/_generar_word.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "AMUCO_Sistema_IA_Harold.docx")

doc = Document()

# ── Márgenes ──────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(3)

# ── Paleta de colores ─────────────────────────────────────
AZUL       = RGBColor(0x1A, 0x3A, 0x5C)   # azul AMUCO
AZUL_LIGHT = RGBColor(0x2A, 0x52, 0x98)
GRIS       = RGBColor(0x55, 0x55, 0x55)
NEGRO      = RGBColor(0x1A, 0x1A, 0x1A)
VERDE      = RGBColor(0x2E, 0x7D, 0x32)
NARANJA    = RGBColor(0xE6, 0x5C, 0x00)

# ── Helpers ───────────────────────────────────────────────

def set_font(run, size=11, bold=False, color=NEGRO, italic=False):
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.color.rgb = color
    run.font.italic = italic
    run.font.name  = "Calibri"

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, size=16, bold=True, color=AZUL)
    # Línea inferior simulada con borde
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '1A3A5C')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=13, bold=True, color=AZUL_LIGHT)
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, size=11, bold=True, color=NEGRO)
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    set_font(run, size=11, color=GRIS)
    return p

def bullet(text, label=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.8)
    if label:
        r1 = p.add_run(f"{label}: ")
        set_font(r1, size=11, bold=True, color=AZUL)
        r2 = p.add_run(text)
        set_font(r2, size=11, color=GRIS)
    else:
        run = p.add_run(text)
        set_font(run, size=11, color=GRIS)
    return p

def nota(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.8)
    run = p.add_run(f"  {text}")
    set_font(run, size=10, italic=True, color=GRIS)
    return p

def etiqueta_estado(texto, color=VERDE):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"  {texto}")
    set_font(run, size=10, bold=True, color=color)
    return p

def separador():
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
r = p.add_run("AMUCO INC.")
set_font(r, size=22, bold=True, color=AZUL)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Sistema de Atención e-Commerce con Inteligencia Artificial")
set_font(r2, size=14, bold=False, color=AZUL_LIGHT)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Línea Coatings & Paints  ·  Marzo 2026")
set_font(r3, size=11, italic=True, color=GRIS)

separador()

# ═══════════════════════════════════════════════════════════
# SECCIÓN 1 — QUÉ ES EL SISTEMA
# ═══════════════════════════════════════════════════════════

heading1("¿Qué es este sistema?")

body("Este documento describe dos herramientas de inteligencia artificial diseñadas específicamente para AMUCO, línea Coatings & Paints. Ambas funcionan de forma independiente y automática — sin que nadie tenga que operar nada manualmente.")

body("En términos simples: son asistentes digitales que trabajan en segundo plano. Uno atiende a los clientes cuando tienen preguntas técnicas. El otro se encarga de hacer seguimiento a las cotizaciones que no han tenido respuesta.")

separador()

# ═══════════════════════════════════════════════════════════
# SECCIÓN 2 — AGENTE 1: CHATBOT
# ═══════════════════════════════════════════════════════════

heading1("Herramienta 1 — Asistente Técnico de Ventas")

heading2("¿Qué hace?")
body("Responde preguntas de clientes sobre los productos de la línea Coatings & Paints — en cualquier momento del día, en el idioma del cliente (español, inglés o portugués), y con información técnica real.")

body("Funciona tanto desde una página web como por WhatsApp. El cliente escribe su pregunta y recibe una respuesta en segundos, sin esperar a que un asesor esté disponible.")

heading2("¿Cómo sabe lo que responde?")
body("El sistema leyó y procesó todas las fichas técnicas, hojas de seguridad y documentos de los 33 productos de la línea — en total 412 archivos. Esa información quedó almacenada en una base de datos interna.")

body("Cuando un cliente hace una pregunta, el sistema busca en esa base de datos, encuentra los documentos más relevantes, los lee, y redacta una respuesta con los datos reales. No inventa nada — si la información no está en los documentos, lo dice.")

heading2("¿Cómo entiende la pregunta?")
body("Antes de buscar, el sistema clasifica la pregunta en dos categorías:")
bullet("Pregunta general — por ejemplo: '¿Qué productos tienen para pinturas?'. En este caso responde desde el catálogo general de productos.")
bullet("Pregunta específica — por ejemplo: '¿Cuál es la viscosidad del PLASTIKOTE?'. En este caso busca la ficha técnica de ese producto y responde con los datos exactos.")
body("Esta clasificación ocurre automáticamente, sin intervención humana y sin costo adicional.")

heading2("¿Recuerda el contexto de la conversación?")
body("Sí. Si un cliente pregunta sobre un producto y luego hace una pregunta de seguimiento ('¿y en qué superficies se aplica?'), el sistema recuerda de qué producto estaban hablando. No es necesario repetir el nombre del producto en cada mensaje.")

separador()

# ═══════════════════════════════════════════════════════════
# SECCIÓN 3 — WHATSAPP
# ═══════════════════════════════════════════════════════════

heading1("WhatsApp — Cómo funciona y qué se necesita")

heading2("Estado actual: entorno de pruebas")
body("Para esta demostración estamos usando una herramienta de desarrollo llamada Twilio Sandbox. Twilio es una empresa internacional que provee infraestructura de comunicaciones — en términos simples, actúa como el puente entre WhatsApp y nuestro sistema.")

body("El Sandbox es un entorno de pruebas gratuito que permite verificar que todo funciona correctamente antes de contratar el servicio oficial. Es el equivalente a una prueba piloto: funciona igual que la versión real, pero con restricciones menores de uso.")

nota("Lo que se ve en esta demo es exactamente lo que recibiría un cliente real. La diferencia es que en producción el mensaje llega desde el número oficial de AMUCO.")

heading2("¿Cómo funciona el flujo por WhatsApp?")
body("El proceso es simple:")
bullet("El cliente escribe al número de WhatsApp de AMUCO.")
bullet("El mensaje llega a nuestro sistema a través de Twilio.")
bullet("El asistente procesa la pregunta y genera la respuesta.")
bullet("La respuesta llega de vuelta al cliente en segundos.")
body("Todo esto ocurre de forma automática. Twilio es solo el canal — el asistente y la inteligencia están del lado de AMUCO.")

heading2("Para ir a producción — tres opciones")
body("Para usar el asistente con el número oficial de AMUCO, hay tres caminos posibles:")

heading3("Opción A — WhatsApp Business (App gratuita)")
body("Es la app que ya existe para negocios pequeños. Permite tener un perfil de empresa, catálogo de productos y respuestas rápidas. No permite automatización — todo se responde manualmente. No es compatible con este sistema.", indent=True)
etiqueta_estado("  No aplicable para automatización", color=NARANJA)

heading3("Opción B — WhatsApp Business API vía Twilio (Recomendada)")
body("Se contrata el servicio de API de WhatsApp a través de Twilio. Permite automatización completa, número propio de AMUCO, sin límite de conversaciones simultáneas. El proceso de aprobación toma entre 1 y 2 semanas.", indent=True)
etiqueta_estado("  Recomendada — compatible con este sistema", color=VERDE)

heading3("Opción C — WhatsApp Business API directo con Meta")
body("Se solicita la API directamente a Meta (empresa dueña de WhatsApp). Mayor control y sin intermediarios, pero el proceso de aprobación es más largo (2 a 4 semanas) y requiere cumplir requisitos de verificación de negocio.", indent=True)
etiqueta_estado("  Compatible — proceso más largo", color=GRIS)

separador()

# ═══════════════════════════════════════════════════════════
# SECCIÓN 4 — AGENTE 2: RECORDATORIOS
# ═══════════════════════════════════════════════════════════

heading1("Herramienta 2 — Seguimiento Automático de Cotizaciones")

heading2("¿Qué hace?")
body("Revisa el sistema de cotizaciones de AMUCO (kiwi.amucoinc.com) y detecta automáticamente qué clientes tienen una oferta pendiente desde hace 2 o más días sin respuesta. Luego envía un email de seguimiento personalizado a cada uno.")

body("El asesor comercial no tiene que recordar hacer el seguimiento ni revisar manualmente cuáles cotizaciones están sin respuesta — el sistema lo hace por él, todos los días.")

heading2("¿Cómo accede al sistema de cotizaciones?")
body("El sistema de AMUCO no tiene una puerta de acceso técnica oficial (lo que en tecnología se llama API). Para resolver esto, mapeamos los mismos movimientos que hace el sistema interno cuando uno navega por el CRM — el mismo proceso que hace el navegador cuando el asesor revisa las cotizaciones.")

body("En términos prácticos: el agente entra al sistema como si fuera un usuario normal, con las credenciales del asesor, revisa las cotizaciones y extrae la información necesaria. No modifica nada — solo lee.")

heading2("¿Qué información usa para redactar el email?")
body("Por cada cotización sin respuesta, el agente obtiene automáticamente:")
bullet("Nombre del cliente y empresa")
bullet("Producto cotizado, cantidad y precio")
bullet("Fecha de la cotización y condiciones (incoterm, ETD)")
bullet("Correos electrónicos del cliente registrados en el sistema")

heading2("¿Cómo detecta el idioma del cliente?")
body("El sistema analiza el nombre de la empresa para determinar si el cliente es hispanohablante, anglohablante o lusohablante, y redacta el email en ese idioma. El email de seguimiento llega al cliente en su propio idioma, sin intervención manual.")

heading2("¿Los emails suenan a Harold?")
body("Sí. Los templates de email fueron entrenados con correos reales del asesor comercial. El estilo, la estructura y el tono son los mismos que Harold usaría — saludo personalizado, referencia al producto y monto específicos, mención de condiciones del mercado como contexto de urgencia, y un cierre directo.")

nota("El cliente recibe un email que parece escrito por Harold — porque está basado en cómo Harold realmente escribe.")

heading2("¿Harold sabe cuándo se envía un recordatorio?")
body("Sí. Cada vez que el agente envía un recordatorio a un cliente, Harold recibe automáticamente una notificación interna por email con el detalle: qué cliente fue contactado, con qué cotización, a qué correo se envió y por qué monto.")

heading2("Modos de operación")
body("El agente tiene tres modos para garantizar que nada se envíe sin control:")
bullet("Modo revisión", "Solo muestra qué cotizaciones tienen seguimiento pendiente y qué habría enviado — sin enviar nada. Útil para revisar.")
bullet("Modo prueba", "Envía el email, pero a una dirección interna en lugar del cliente real. Permite verificar el contenido antes de activar.")
bullet("Modo producción", "Envía directamente al cliente. Requiere aprobación explícita para activar.")

separador()

# ═══════════════════════════════════════════════════════════
# SECCIÓN 5 — ESTADO ACTUAL
# ═══════════════════════════════════════════════════════════

heading1("Estado actual del proyecto")

body("El sistema está en fase de desarrollo avanzado y validación. Ambas herramientas funcionan en entorno controlado y han sido probadas con datos reales de AMUCO.")

table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'

# Encabezados
hdr = table.rows[0].cells
hdr[0].text = "Componente"
hdr[1].text = "Estado"
hdr[2].text = "Notas"

for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), '1A3A5C')
        cell._tc.get_or_add_tcPr().append(shading)

filas = [
    ("Chatbot web",                  "Funcionando",        "Disponible en localhost:5001"),
    ("Chatbot por WhatsApp",         "Funcionando",        "Twilio Sandbox activo"),
    ("Agente de recordatorios",      "En pruebas",         "Modo prueba — envíos a email interno"),
    ("WhatsApp número propio AMUCO", "Pendiente",          "Requiere aprobación Meta"),
]

for i, (comp, estado, nota_txt) in enumerate(filas, start=1):
    row = table.rows[i].cells
    row[0].text = comp
    row[1].text = estado
    row[2].text = nota_txt
    for cell in row:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

separador()

# ═══════════════════════════════════════════════════════════
# SECCIÓN 6 — PRÓXIMOS PASOS
# ═══════════════════════════════════════════════════════════

heading1("Próximos pasos para ir a producción")

body("Para activar el sistema completamente, se requieren dos aprobaciones:")

heading3("1. WhatsApp con número propio de AMUCO")
body("Registrar un número de negocio de AMUCO a través de la API de WhatsApp Business (Opción B recomendada). El proceso toma entre 1 y 2 semanas e implica verificar el negocio ante Meta.", indent=True)

heading3("2. Activación del agente de recordatorios en producción")
body("Revisar 2-3 emails generados automáticamente, confirmar que el tono y contenido es el correcto, y dar aprobación para que el agente envíe directamente a los clientes reales.", indent=True)

separador()

# ═══════════════════════════════════════════════════════════
# SECCIÓN 7 — COSTOS
# ═══════════════════════════════════════════════════════════

heading1("Costos operativos estimados")

body("Una vez en producción, los costos de operación son bajos y predecibles:")

bullet("Modelo de IA (Google Gemini)", "aprox. $0.002 por conversación del chatbot. Con uso normal, menos de $5 al mes.")
bullet("WhatsApp Business API (Twilio)", "aprox. $0.005 por mensaje. Con 200 conversaciones al mes, menos de $15.")
bullet("Infraestructura", "el sistema corre localmente — sin costos de servidor en la nube.")

nota("Los costos son en USD y pueden variar según el volumen de uso. Son estimaciones basadas en tarifas actuales de los proveedores.")

separador()

# ═══════════════════════════════════════════════════════════
# PIE DE PÁGINA
# ═══════════════════════════════════════════════════════════

p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_footer.paragraph_format.space_before = Pt(30)
r_footer = p_footer.add_run("Documento preparado por Jarvis — Sistema Keystone  ·  Marzo 2026  ·  Confidencial")
set_font(r_footer, size=9, italic=True, color=GRIS)

# ── Guardar ───────────────────────────────────────────────
doc.save(OUTPUT)
print(f"Documento generado: {OUTPUT}")
